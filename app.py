from cs50 import SQL
from flask import Flask, redirect, render_template, request, session, jsonify, send_file, Response, flash
from flask_session import Session
from werkzeug.security import check_password_hash, generate_password_hash
from datetime import timedelta, datetime, date
from helpers import login_required, add_app, check_duplicate
from email_validator import validate_email, EmailNotValidError
import re
from werkzeug.utils import secure_filename
from docx import Document
from collections import defaultdict
import os
from zipfile import ZipFile

# Configuration part copied from CS50 Finance pset
# Configure application
app = Flask(__name__)

#secret key
app.secret_key = b'\x07)\x11\x08\xc8Z|u`*\xefG#\x88>J\xb7\x12\xbe\x99[\xf7\x88+'

#configure session so users stay logged in for the duration of the browser
app.config["SESSION_PERMANENT"] = False
app.config["SESSION_TYPE"] = "filesystem"
Session(app)

@app.after_request
def after_request(response):
    """Ensure responses aren't cached"""
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Expires"] = 0
    response.headers["Pragma"] = "no-cache"
    return response

# configure database connection
db = SQL("sqlite:///turfschuur.db")
db.execute("PRAGMA foreign_keys = ON")

# home
@app.route("/")
def index():
    return render_template("index.html")

# login
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "GET":
        # check if user is already logged in
        if session.get("user_id") is not None:
            return redirect("/dashboard")
        return render_template("login.html")
    else:
        if not request.form.get("name"):
            return render_template("apology.html", apology="Naam is niet ingevuld"), 400
        if not request.form.get("password"):
            return render_template("apology.html", apology="Wachtwoord is niet ingevuld"), 400
        
        # query database for name
        rows = db.execute(
            "SELECT * FROM bestuursleden WHERE naam = ?", request.form.get("name")
        )

        # check if name and password are correct and if he/she is an approved board member
        if len(rows) != 1 or not check_password_hash(rows[0]["hash"], request.form.get("password")) or rows[0]["toegelaten"] < 1:
            return render_template("apology.html", apology="Inloggen mislukt"), 400
        
        # if everything is correct, store id inside session
        session["user_id"] = rows[0]["bestuurslid_id"]
        return redirect("/dashboard")
    
# Logout
@app.route("/logout", methods=["GET"])
@login_required
def logout():
    session.clear()
    return redirect("/")
    
# register for an account
@app.route("/registreren", methods=["GET", "POST"])
def register():
    if request.method == "GET":
        return render_template("registreren.html")
    else:
        # TODO
        if not request.form.get("name"):
            return render_template("apology.html", apology="Geen naam ingevuld"), 400
        if not request.form.get("email"):
            return render_template("apology.html", apology="E-mailadres niet ingevuld"), 400
        # Check if the email is valid
        try:
            v = validate_email(request.form.get("email"))
            # use normalized form of emailadress
            email = v["email"]
        except EmailNotValidError as e:
            return render_template("apology.html", apology="E-mailadres wordt niet herkend"), 400
        # check if password and password-repeat are the same
        if not request.form.get("password") == request.form.get("password_repeat"):
            return render_template("apology.html", apology="Wachtwoorden komen niet overeen"), 400
        
        # update database with new user
        # "toegelaten" must be set to zero (also happens automatically if no value is specified), as you dont want anyone being able to register
        try:
            db.execute(
                """INSERT INTO bestuursleden
                (naam, email, hash, toegelaten)
                VALUES (?, ?, ?, ?)""",
                request.form.get("name"), request.form.get("email"), generate_password_hash(request.form.get("password")), 0
            )
        except:
            return render_template("apology.html", apology="Naam of e-mailadres is al in gebruik"), 400
        # If user is registered, show a page with explaination that another member has to accept him/her
        return render_template("succes.html")
    
# dashboard for members
@app.route("/dashboard", methods=["GET", "POST"])
@login_required
def dashboard():
    if request.method == "GET":
        return render_template("dashboard.html")
    # post for posting appointments
    else:
        # Check if the form has an input field called flag
        flag = request.form.get("updatedelete-flag")
        if flag is not None:
            # Handle the update/delete appointment form
            # Update alle = 0
            # Update enkele = 1
            # Verwijder alle = 2
            # Verwijder enkele = 3

            # Get all info
            id = request.form.get("id")
            titel = request.form.get("app-titel")
            begin = datetime.strptime(request.form.get("app-begin"), "%Y-%m-%dT%H:%M")
            eind = datetime.strptime(request.form.get("app-eind"), "%Y-%m-%dT%H:%M")
            prijs = request.form.get("app-prijs")
            prijs_omschrijving = request.form.get("app-omschrijving")
            extra = request.form.get("app-extra")
            extra_omschrijving = request.form.get("app-omschrijving_extra")
            info = request.form.get("app-info")
            # Check if omschrijving_prijs en extra are empty
            if prijs_omschrijving == "":
                prijs_omschrijving = None
            if extra_omschrijving == "":
                extra_omschrijving = None
            if not titel or not begin or not eind:
                return render_template("apology.html", apology="Afspraak moet minstens een titel, begin- en eindtijd bevatten"), 400
            if eind <= begin:
                return render_template("apology.html", apology="Eindtijd moet later dan begintijd zijn"), 400
            if ((eind - begin) >= timedelta(days=1)):
                return render_template("apology.html", apology="Afspraak duurt te lang"), 400
            
            # Switch case for the different cases like updating/deleting and all/single
            match int(flag):
                case 0:
                    # Update all
                    # Select all targeted appointments
                    appointments = db.execute("""SELECT afspraak_id, begin, eind
                                              FROM afspraken
                                              WHERE reeks_id = (
                                              SELECT reeks_id
                                              FROM afspraken
                                              WHERE afspraak_id = ?
                                              )""", id)
                    if len(appointments) <= 0:
                        appointments = db.execute("""SELECT afspraak_id, begin, eind
                                                  FROM afspraken
                                                  WHERE afspraak_id = ?""", 
                                                  id)
                    
                    # Calculate time difference
                    # Get correct appointment for comparing
                    right_app_time = db.execute("""SELECT begin, eind
                                                FROM afspraken WHERE afspraak_id = ?""", id)
                    original_begin = datetime.fromisoformat(right_app_time[0]["begin"])
                    original_end = datetime.fromisoformat(right_app_time[0]["eind"])
                    timediff_begin = begin - original_begin
                    timediff_eind = eind - original_end
                    # For every appointment in the series, update everything
                    for app in appointments:
                        # Calculate the new datetimes
                        new_begin = datetime.fromisoformat(app["begin"]) + timediff_begin
                        new_eind = datetime.fromisoformat(app["eind"]) + timediff_eind
                        db.execute("""UPDATE afspraken
                                   SET titel = ?, begin = ?, eind = ?, prijs = ?, omschrijving_prijs = ?, extra = ?, omschrijving_extra = ?, info = ?
                                   WHERE afspraak_id = ?""", 
                                   titel, new_begin.isoformat(), new_eind.isoformat(), prijs, prijs_omschrijving, extra, extra_omschrijving, info, app["afspraak_id"])
                case 1:
                    # Update single
                    db.execute("""UPDATE afspraken
                               SET titel = ?, begin = ?, eind = ?, prijs = ?, omschrijving_prijs = ?, extra = ?, omschrijving_extra = ?, info = ?
                               WHERE afspraak_id = ?""", 
                               titel, begin.isoformat(), eind.isoformat(), prijs, prijs_omschrijving, extra, extra_omschrijving, info, id)
                case 2:
                    # Delete all
                    n = db.execute("""DELETE FROM afspraken
                               WHERE reeks_id = (
                               SELECT reeks_id
                               FROM afspraken
                               WHERE afspraak_id = ?
                               )""", 
                               id)
                    if n <= 0:
                        n = db.execute("""DELETE FROM afspraken
                                    WHERE afspraak_id = ?""", 
                                    id)
                    if n <= 0:
                        return render_template("apology.html", apology="Gelieve niet aan de code lopen prutsen"), 400
                case 3:
                    # Delete single
                    db.execute("""DELETE FROM afspraken
                               WHERE afspraak_id = ?""", 
                               id)
                case _:
                    return render_template("apology.html", apology="Gelieve niet aan de code lopen prutsen"), 400

            return redirect("/dashboard")

        if not request.form.get("titel") or not request.form.get("begin") or not request.form.get("einde"):
            return render_template("apology.html", apology="Minstens een titel, begin- en eindtijd instellen"), 400
        # check if end time is later than begin time
        begintijd = datetime.strptime(request.form.get("begin"), "%Y-%m-%dT%H:%M")
        eindtijd = datetime.strptime(request.form.get("einde"), "%Y-%m-%dT%H:%M")
        if eindtijd <= begintijd:
            return render_template("apology.html", apology="Eindtijd moet later dan begintijd zijn"), 400
        if ((eindtijd - begintijd) >= timedelta(days=1)):
            return render_template("apology.html", apology="Afspraak duurt te lang"), 400
        
        titel = request.form.get("titel").strip()
        prijs = float(request.form.get("prijs"))
        desc_prijs = request.form.get("omschrijving")
        extra = float(request.form.get("extra"))
        desc_extra = request.form.get("omschrijving_extra")
        contact = request.form.get("contactgegevens").strip()
        info = request.form.get("info").strip()
        
        vl_raw = request.form.get("voorletter")
        if vl_raw is None:
            vl = None
        else:
            # Get all of the letters
            letters = re.findall("[a-zA-Z]", vl_raw)
            # make everything uppercase and add a dot after every letter
            vl = ((".".join(letters)) + ".").upper()
        an = request.form.get("achternaam").upper().strip()
        
        straat = request.form.get("straat").lower().strip()
        if not straat:
            straat = None
        hn = request.form.get("huisnummer").lower().split()
        huisnummer = "".join(hn)
        if not huisnummer:
            huisnummer = None
        post = request.form.get("postcode").lower().split()
        postcode = "".join(post)
        if not postcode:
            postcode = None
        plaats = request.form.get("woonplaats").lower().strip()
        if not plaats:
            plaats = None
        land = request.form.get("land").strip()
        if not land:
            land = "Nederland"

        repeat = request.form.get("interval")
        repeatx = request.form.get("reeks")

        # If address is filled in, check if the address/person is yet in the database
        if (an or vl) and postcode and huisnummer:
            # Check if initial or surname is empty
            if not an or not vl:
                return render_template("apology.html", apology="Moet zowel een voorletter als een achternaam zijn opgegeven om in het adressenbestand te zoeken"), 400
            # Querie database for right address id
            adres_id = db.execute("""SELECT adres_id
                                    FROM adressenbestand
                                    WHERE voorletter = ? AND achternaam = ? AND postcode = ? AND huisnummer = ?""",
                                    vl, an, postcode, huisnummer)
            # If not, add new address/person to address table (fields cant be null)
            if len(adres_id) <= 0:
                try:
                    id = db.execute("""INSERT INTO adressenbestand
                                    (voorletter, achternaam, straat, huisnummer, woonplaats, land, postcode, contact)
                                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)""", 
                                    vl, an, straat, huisnummer, plaats, land, postcode, request.form.get("contactgegevens"))
                except ValueError:
                    return render_template("apology.html", apology="Niet alle benodigde gegevens zijn ingevuld"), 400
                
                try:
                    add_app(titel, begintijd, eindtijd, prijs, desc_prijs, extra, desc_extra, info, id, repeat, repeatx)
                except ValueError:
                    return render_template("apology.html", apology="Ongeldige herhaling ingesteld"), 400
            # Else if there's a name/address, add appointment with adres_id set to name/address
            else:
                try:
                    add_app(titel, begintijd, eindtijd, prijs, desc_prijs, extra, desc_extra, info, adres_id[0]["adres_id"], repeat, repeatx)
                except ValueError:
                    return render_template("apology.html", apology="Ongeldige herhaling ingesteld"), 400
                
        # Else if address isnt filled in, just add appointment to database with the adres_id set to null
        else:
            try:
                add_app(titel, begintijd, eindtijd, prijs, desc_prijs, extra, desc_extra, info, None, repeat, repeatx)
            except ValueError:
                    return render_template("apology.html", apology="Ongeldige herhaling ingesteld"), 400
        return redirect("/dashboard")
    
# Route for generating invoices
@app.route("/factuur", methods=["GET", "POST"])
@login_required
def factuur():
    if request.method == "GET":
        return render_template("factuur.html")
    else:
        month_string = request.form.get("month")
        if not month_string:
            return render_template("apology.html", apology="Geen maand ingevuld"), 400
        month = month_string.split('-')
        file = request.files['template']
        if file.filename != '':
            # Check if the file is allowed
            x = file.filename.rsplit('.', 1)
            if file and x[1].lower() == "docx":
                # Save file on system
                filename = secure_filename("template.docx")
                file.save('facturen/' + filename)
            else:
                return render_template("apology.html", apology="Gelieve enkel .docx bestanden up te loaden"), 400
            
        # Get all appointments in the specified month
        appointments = db.execute("""SELECT afspraken.*, adressenbestand.*
                                  FROM afspraken
                                  LEFT JOIN adressenbestand ON afspraken.adres_id = adressenbestand.adres_id
                                  WHERE strftime('%m', eind) = ? AND strftime('%Y', eind) = ? AND afspraken.adres_id IS NOT NULL
                                  ORDER BY adres_id""",
                                  month[1], month[0])
        
        if len(appointments) <= 0:
            return "Geen afspraken op naam in deze maand"
        
        # Remove all previous files
        for root, dir, files in os.walk("./facturen/generated"):
            for file in files:
                os.remove(os.path.join(root, file))
        # Make a list of dicts with the keywords to be replaced
        customers = defaultdict(list)
        for app in appointments:
            customers[app['adres_id']].append(app)
        # For each customer, open the doc and replace all keywords
        for adres_id, appt_list in customers.items():
            doc = Document("facturen/template.docx")
            # Replace all keywords for address in the paragraphs
            postcode = appt_list[0]["postcode"]
            pc = re.sub("[A-Za-z]+", lambda ele: " " + ele[0] + " ", postcode).upper()
            total = 0
            for app in appt_list:
                total += float(app["prijs"])
                if app["extra"] is not None:
                    total += float(app["extra"])
            # At end of loop round total to two decimals
            total = "{:.2f}".format(float(total))
            today = date.today()
            maxf = db.execute("""SELECT max(factuurnummer) AS max FROM facturen""")
            if len(maxf) <= 0 or maxf[0]["max"] is None:
                factuurnummer = 50241
            else:
                factuurnummer = int(maxf[0]["max"]) + 1
            adres = {
                "{{Voorletters}}": appt_list[0]["voorletter"],
                "{{Achternaam}}": appt_list[0]["achternaam"],
                "{{Straat}}": str(appt_list[0]["straat"]).upper(),
                "{{Huisnummer}}": appt_list[0]["huisnummer"],
                "{{Postcode}}": pc,
                "{{Woonplaats}}": str(appt_list[0]["woonplaats"]).upper(),
                "{{Totaal}}": "€" + str(total),
                "{{Factuurdatum}}": today.strftime("%d-%m-%Y"),
                "{{Factuurnummer}}": str(factuurnummer),
                }
            
            for paragraph in list(doc.paragraphs):
                for old, new in adres.items():
                    if old in paragraph.text:
                        paragraph.text = paragraph.text.replace(old, new)
            
            # put costs in table
            # Search for the correct table
            table = None
            for t in list(doc.tables):
                if len(t.columns) == 3:
                    table = t
            if table is None:
                return render_template("apology.html", apology="geen tabel gevonden"), 400
            
            for app in appt_list:
                row = table.add_row().cells
                omsch_pr = app["omschrijving_prijs"]
                if omsch_pr is None:
                    omsch_pr = ""
                row[0].text = omsch_pr
                pr = "{:.2f}".format(float(app["prijs"]))
                row[1].text = str("€" + pr)
                row[2].text = "0,00%"
                if app["omschrijving_extra"] is not None and app["extra"] is not None:
                    row2 = table.add_row().cells
                    row2[0].text = app["omschrijving_extra"]
                    xpr = "{:.2f}".format(float(app["extra"]))
                    row2[1].text = "€" + xpr
                    row2[2].text = "0,00%"

            savefilename = "facturen/generated/" + appt_list[0]["achternaam"] + str(adres_id) + "_" + today.strftime("%d%m%Y") + ".docx"
            doc.save(savefilename)

            # Update database
            try:
                db.execute("""INSERT INTO facturen (factuurnummer, adres_id)
                           VALUES (?, ?)""", 
                           factuurnummer, adres_id)
            except:
                return render_template("apology.html", apology="Fout met genereren van facturen, probeer het later opnieuw"), 500
        # Zip all the generated invoices and send it back to the user
        with ZipFile("./facturen/facturen.zip", "w") as zip:
            for root, dir, files in os.walk("./facturen/generated"):
                for file in files:
                    zip.write(os.path.join(root, file), arcname=file)

        return send_file("facturen/facturen.zip")
    
# Address database
@app.route("/adressenbestand", methods=["GET"])
@login_required
def adressenbestand():
    adressen = db.execute("""SELECT *
                            FROM adressenbestand""")
    return render_template("adressenbestand.html", adressen=adressen)

# Board members
@app.route("/bestuursleden", methods=["GET"])
@login_required
def bestuursleden():
    bestuursleden = db.execute("""SELECT bestuurslid_id, naam, email
                               FROM bestuursleden
                               WHERE toegelaten = 1""")
    
    nnt = db.execute("""SELECT bestuurslid_id, naam, email
                     FROM bestuursleden
                     WHERE toegelaten = 0""")
    return render_template("bestuursleden.html", bestuursleden=bestuursleden, nnt=nnt)

# Route for ajax calendar request
@app.route("/app_month", methods=["GET"])
@login_required
def get_appointments_month():
    # Get the month of which the appointments should be loaded
    month = request.args.get("month")
    year = request.args.get("year")
    mstring = str(month).zfill(2)
    myear = str(year)
    months = db.execute("""SELECT afspraak_id, begin, eind FROM afspraken
                        WHERE strftime('%m', begin) = ? AND strftime('%Y', begin) = ?
                        ORDER BY begin ASC""",
                        mstring, myear)
    return jsonify(months)

# Route for ajax request for appointments on a day
@app.route("/app_day", methods=["GET"])
@login_required
def get_appointments_day():
    # Get the day
    day_str = request.args.get("day")
    # If something went wrong with the ajax request, just return an empty list
    if day_str is None:
        jsonify([])
    day = datetime.strptime(day_str, '%a %b %d %Y')
    
    # Get end of day (or start new day)
    endday = day + timedelta(days=1)
    # Convert both back to iso 8601 format
    day.replace(microsecond=0).isoformat()
    endday.replace(microsecond=0).isoformat()
    # Query database for every appointment where the starttime <= endday and endtime >= day
    day_appointments = db.execute("""SELECT afspraken.*, adressenbestand.adres_id, adressenbestand.voorletter, adressenbestand.achternaam
                                  FROM afspraken
                                  LEFT JOIN adressenbestand ON afspraken.adres_id = adressenbestand.adres_id
                                  WHERE begin <= ? AND eind >= ?""",
                                  endday, day)
    
    return jsonify(day_appointments)

# Route for deleting address entries
@app.route("/delete_address", methods=["POST"])
@login_required
def delete_address():
    adres_id = request.form.get("adres_id")
    # Delete all appointments linked to this address
    db.execute("""DELETE FROM afspraken
               WHERE adres_id = ?""",
               adres_id)
    
    # Delete address
    db.execute("""DELETE FROM adressenbestand
               WHERE adres_id = ?""",
               adres_id)
    return Response(status=204)

# Route for deleting board member
@app.route("/delete_member", methods=["POST"])
@login_required
def delete_member():
    db.execute("""DELETE FROM bestuursleden
                WHERE bestuurslid_id = ?""",
                request.form.get("bestuurslid_id"))
    return Response(status=200)
    
@app.route("/accept_member", methods=["POST"])
@login_required
def accept_member():
    db.execute("""UPDATE bestuursleden
               SET toegelaten = 1
               WHERE bestuurslid_id = ?""",
               request.form.get("bestuurslid_id"))
    return Response(status=200)

# Route for checking for duplicate appointments
@app.route("/check_duplicate", methods=["GET"])
@login_required
def check_duplicate_apps():
    b = datetime.strptime(request.args.get("begin"), "%Y-%m-%dT%H:%M")
    e = datetime.strptime(request.args.get("eind"), "%Y-%m-%dT%H:%M")
    print(f"checking: {b} and {e}")

    # Call function for checking duplicate
    return jsonify(check_duplicate(b, e))